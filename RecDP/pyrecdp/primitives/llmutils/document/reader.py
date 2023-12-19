import os
from abc import abstractmethod, ABC
from pathlib import Path
from typing import List, Optional, Dict, Type
from pyrecdp.core.import_utils import check_availability_and_install
from pyrecdp.primitives.llmutils.document.schema import Document

class DocumentReader(ABC):
    """interface for document loader"""

    @abstractmethod
    def load(self) -> List[Document]:
        """abstract method for load documents"""


class FileBaseReader(DocumentReader, ABC):
    """interface for loading document from a file."""

    def __init__(self, file: Path, single_text_per_document: bool = True, page_separator: str = '\n'):
        self.single_text_per_document = single_text_per_document
        self.page_separator = page_separator or '\n'
        self.file = file
        
    @classmethod
    def setup(cls):
        for pkg in cls.system_requirements:
            os.system(f'apt-get install -y {pkg}')
        for pkg in cls.requirements:
            check_availability_and_install(pkg, verbose=1)

    def get_metadata(self):
        return {"source": str(self.file)}

    def load(self) -> List[Document]:
        docs: List[Document] = self.load_file(self.file)
        docs: List[Document] = list(filter(lambda d: (d.text.strip() != ""), docs))

        def firstAlphaIsUppercase(word: str) -> bool:
            i: int = 0
            while i < len(word):
                char = doc.text[i]
                if char.isalpha():
                    return char.isupper()
                i += 1
            return False

        if self.single_text_per_document:
            text = ''
            for doc in docs:
                if firstAlphaIsUppercase(doc.text):
                    text += f"\n\n {doc.text}"
                else:
                    text += f" {doc.text}"

            return [Document(text=text, metadata=self.get_metadata())]
        else:
            return docs

    @abstractmethod
    def load_file(self, file: Path) -> List[Document]:
        """Load data from the input directory."""


class PDFReader(FileBaseReader):
    """PDF parser."""
    system_requirements = []
    requirements = ['pypdf']
    def __init__(self, file: Path, single_text_per_document: bool = True, page_separator: str = '\n',
                 **load_kwargs):
        super().__init__(file, single_text_per_document, page_separator)
        self.load_kwargs = load_kwargs
        self.file = file

    def load_file(self, file: Path) -> List[Document]:
        import pypdf
        # Create a PDF object
        pdf = pypdf.PdfReader(file, **self.load_kwargs)

        # Get the number of pages in the PDF document
        num_pages = len(pdf.pages)

        # Iterate over every page
        docs = []
        for page in range(num_pages):
            # Extract the text from the page
            page_text = pdf.pages[page].extract_text()
            page_label = pdf.page_labels[page]
            metadata = {"page_label": page_label, "source": str(file)}
            docs.append(Document(text=page_text, metadata=metadata))

        return docs


class DocxReader(FileBaseReader):
    """Docx parser."""
    system_requirements = []
    requirements = ['python-docx']
    def __init__(self, file: Path, single_text_per_document: bool = True, page_separator: str = '\n'):
        super().__init__(file, single_text_per_document, page_separator)

    def load_file(self, file: Path) -> List[Document]:
        """Parse file."""
        import docx
        document = docx.Document(str(file))

        # read in each paragraph in file
        return [Document(text=p.text, metadata={"source": str(file)}) for p in document.paragraphs]


class ImageReader(FileBaseReader):
    """Image parser.

    Extract text from images using pytesseract.

    """
    system_requirements = ['tesseract-ocr']
    requirements = ['pillow', 'pytesseract']
    def __init__(
            self,
            file: Path,
            single_text_per_document: bool = True,
            page_separator: str = '\n',
            keep_image: bool = False,
    ):
        super().__init__(file, single_text_per_document, page_separator)
        self._keep_image = keep_image

    def load_file(self, file: Path) -> List[Document]:
        """Parse file."""
        from PIL import Image
        from pytesseract import pytesseract
        # load document image
        image = Image.open(file)
        if image.mode != "RGB":
            image = image.convert("RGB")

        # Parse image into text
        text_str = pytesseract.image_to_string(image)

        return [
            Document(text=text_str, metadata={"source": str(file)})
        ]


class AudioReader(FileBaseReader):
    system_requirements = ['ffmpeg']
    requirements = ['openai-whisper']
    def __init__(
            self,
            file: Path,
            single_text_per_document: bool = True,
            page_separator: str = '\n',
            model = 'small',
    ):
        super().__init__(file, single_text_per_document, page_separator)
        import whisper
        self.model = whisper.load_model(model)

    def transcribe(self, file):
        return self.model.transcribe(file)
    def load_file(self, file: Path) -> List[Document]:
        file = str(file)
        result = self.transcribe(file)
        return [
            Document(text=result['text'], metadata={"source": file, 'language': result['language']})
        ]


CUSTOMIZE_SUPPORTED_SUFFIX: Dict[str, Type[FileBaseReader]] = {
    ".pdf": PDFReader,
    ".docx": DocxReader,
    ".jpg": ImageReader,
    ".jpeg": ImageReader,
    ".png": ImageReader,
    ".mp3": AudioReader,
    ".wav": AudioReader,
    ".flac": AudioReader,
}


class DirectoryReader(DocumentReader):
    def __init__(
            self,
            input_dir: Optional[str] = None,
            glob: str = "**/[!.]*",
            use_multithreading: bool = False,
            max_concurrency: Optional[int] = None,
            input_files: Optional[List] = None,
            single_text_per_document: bool = True,
            exclude: Optional[List] = None,
            exclude_hidden: bool = True,
            silent_errors: bool = False,
            recursive: bool = False,
            encoding: str = "utf-8",
            required_exts: Optional[List[str]] = CUSTOMIZE_SUPPORTED_SUFFIX.keys(),
            page_separator: Optional[str] = '\n',
    ) -> None:
        """
       Loads documents from a directory or a list of files.

       Args:
           input_dir: The input directory.
           glob: A glob pattern to match files.
           recursive: Whether to recursively search the input directory.
           use_multithreading: Whether to use multithreading to load documents.
           max_concurrency: The maximum number of concurrent threads to use.
           input_files: A list of input files.
           single_text_per_document: Whether to load each file as a single document.
           exclude: A list of file patterns to exclude from loading.
           exclude_hidden: Whether to exclude hidden files from loading.
           silent_errors: Whether to silently ignore errors when loading documents.
           encoding: The encoding to use when loading documents.
           required_exts: A list of file extensions that are required for documents.
                          default extensions are [.pdf, .docx, .jpeg, .jpg, .png]
       """
        if not input_dir and not input_files:
            raise ValueError("Must provide either `path` or `input_files`.")
        self.glob = glob
        self.use_multithreading = use_multithreading
        self.max_concurrency = max_concurrency
        self.encoding = encoding
        self.silent_errors = silent_errors
        self.exclude = exclude
        self.recursive = recursive
        self.exclude_hidden = exclude_hidden
        self.required_exts = required_exts
        self.page_separator = page_separator
        self.file_extractor = {}
        if input_files:
            self.input_files = []
            for path in input_files:
                if not os.path.isfile(path):
                    raise ValueError(f"File {path} does not exist.")
                input_file = Path(path)
                self.input_files.append(input_file)
        elif input_dir:
            if not os.path.isdir(input_dir):
                raise ValueError(f"Directory {input_dir} does not exist.")
            self.input_dir = Path(input_dir)
            self.exclude = exclude
            self.input_files = self._add_files(self.input_dir)

        if len(self.input_files) == 1:
            self.use_multithreading = False

        self.single_text_per_document = single_text_per_document
        
    def setup(self):
        suffix_list = set(input_file.suffix.lower() for input_file in self.input_files)
        for file_suffix in suffix_list:
            if file_suffix in CUSTOMIZE_SUPPORTED_SUFFIX:
                if file_suffix not in self.file_extractor:
                    file_base_reader_cls: Type[FileBaseReader] = CUSTOMIZE_SUPPORTED_SUFFIX[file_suffix]
                    file_base_reader_cls.setup()

    def _add_files(self, input_dir: Path) -> List[Path]:
        all_files = set()
        rejected_files = set()

        if self.exclude is not None:
            for excluded_pattern in self.exclude:
                if self.recursive:
                    # Recursive glob
                    for file in input_dir.rglob(excluded_pattern):
                        rejected_files.add(Path(file))
                else:
                    # Non-recursive glob
                    for file in input_dir.glob(excluded_pattern):
                        rejected_files.add(Path(file))

        p = Path(input_dir)
        file_refs = list(p.rglob(self.glob) if self.recursive else p.glob(self.glob))

        for ref in file_refs:
            # Manually check if file is hidden or directory instead of
            # in glob for backwards compatibility.
            is_dir = ref.is_dir()
            skip_because_hidden = self.exclude_hidden and ref.name.startswith(".")
            skip_because_bad_ext = (
                    self.required_exts is not None and ref.suffix not in self.required_exts
            )
            skip_because_excluded = ref in rejected_files

            if (
                    is_dir
                    or skip_because_hidden
                    or skip_because_bad_ext
                    or skip_because_excluded
            ):
                continue
            else:
                all_files.add(ref)

        new_input_files = sorted(all_files)

        if len(new_input_files) == 0:
            raise ValueError(f"No files found in {input_dir}.")

        return new_input_files

    def _load_file(self, input_file: Path, pbar):
        try:
            file_suffix = input_file.suffix.lower()
            if file_suffix in CUSTOMIZE_SUPPORTED_SUFFIX:                    
                if file_suffix not in self.file_extractor:
                    file_base_reader_cls: Type[FileBaseReader] = CUSTOMIZE_SUPPORTED_SUFFIX[file_suffix]
                    self.file_extractor[file_suffix] = file_base_reader_cls(
                        input_file,
                        single_text_per_document=self.single_text_per_document,
                        page_separator=self.page_separator,
                    )
                loader = self.file_extractor[file_suffix]
                return loader.load()
            else:
                from pyrecdp.core.import_utils import import_langchain
                import_langchain()
                from langchain.document_loaders import UnstructuredFileLoader
                loader = UnstructuredFileLoader(str(input_file))
                docs = [Document(text=doc.text, metadata=doc.metadata) for doc in loader.load()]
                docs = list(filter(lambda d: (d.pa.strip() != ""), docs))
                if self.single_text_per_document:
                    text = self.page_separator.join([doc.text for doc in docs])
                    return [Document(text=text, metadata={"source": str(input_file)})]
                else:
                    return docs
        finally:
            if pbar:
                pbar.update(1)

    def load(self) -> List[Document]:
        from tqdm import tqdm
        pbar = tqdm(total=len(self.input_files))
        try:
            docs_result: List[Document] = []
            if self.use_multithreading:
                from concurrent.futures import ThreadPoolExecutor
                with ThreadPoolExecutor(self.max_concurrency) as executor:
                    for docs in executor.map(lambda i: self._load_file(i, pbar), self.input_files):
                        docs_result.extend(docs)
            else:
                for file in self.input_files:
                    docs = self._load_file(file, pbar)
                    if len(docs) > 0:
                        docs_result.extend(docs)
            return docs_result
        finally:
            pbar.close()
