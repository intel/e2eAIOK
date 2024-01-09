"""
 Copyright 2024 Intel Corporation

 Licensed under the Apache License, Version 2.0 (the "License");
 you may not use this file except in compliance with the License.
 You may obtain a copy of the License at

      https://www.apache.org/licenses/LICENSE-2.0

 Unless required by applicable law or agreed to in writing, software
 distributed under the License is distributed on an "AS IS" BASIS,
 WITHOUT WARRANTIES OR CONDITIONS OF ANY KIND, either express or implied.
 See the License for the specific language governing permissions and
 limitations under the License.
 """

import os
import re
from abc import abstractmethod, ABC
from pathlib import Path
from typing import List, Dict, Type, Optional, Callable, Any, Union, Sequence

from loguru import logger

from pyrecdp.core.import_utils import check_availability_and_install
from pyrecdp.primitives.document.schema import Document


class DocumentReader(ABC):
    """interface for document loader"""

    @abstractmethod
    def load(self) -> List[Document]:
        """abstract method for load documents"""


class FileBaseReader(DocumentReader, ABC):
    """interface for loading document from a file."""

    def __init__(self, file: Path, single_text_per_document: bool = True, page_separator: str = '\n'):
        self.single_text_per_document = single_text_per_document
        self.page_separator = page_separator or '\n'
        self.file = file

    @classmethod
    def setup(cls):
        for pkg in cls.system_requirements:
            os.system(f'apt-get install -y {pkg}')
        for pkg in cls.requirements:
            check_availability_and_install(pkg, verbose=1)

    def get_metadata(self):
        return {"source": str(self.file)}

    def load(self) -> List[Document]:
        docs: List[Document] = self.load_file(self.file)
        docs: List[Document] = list(filter(lambda d: (d.text.strip() != ""), docs))

        def firstAlphaIsUppercase(word: str) -> bool:
            i: int = 0
            while i < len(word):
                char = doc.text[i]
                if char.isalpha():
                    return char.isupper()
                i += 1
            return False

        if self.single_text_per_document:
            text = ''
            for doc in docs:
                if firstAlphaIsUppercase(doc.text):
                    text += f"\n\n {doc.text}"
                else:
                    text += f" {doc.text}"

            return [Document(text=text, metadata=self.get_metadata())]
        else:
            return docs

    @abstractmethod
    def load_file(self, file: Path) -> List[Document]:
        """Load data from the input directory."""


class PDFReader(FileBaseReader):
    """PDF parser."""
    system_requirements = []
    requirements = ['pypdf']

    def __init__(self, file: Path, single_text_per_document: bool = True, page_separator: str = '\n',
                 **load_kwargs):
        super().__init__(file, single_text_per_document, page_separator)
        self.load_kwargs = load_kwargs
        self.file = file

    def load_file(self, file: Path) -> List[Document]:
        import pypdf
        # Create a PDF object
        pdf = pypdf.PdfReader(file, **self.load_kwargs)

        # Get the number of pages in the PDF document
        num_pages = len(pdf.pages)

        # Iterate over every page
        docs = []
        for page in range(num_pages):
            # Extract the text from the page
            page_text = pdf.pages[page].extract_text()
            page_label = pdf.page_labels[page]
            metadata = {"page_label": page_label, "source": str(file)}
            docs.append(Document(text=page_text, metadata=metadata))

        return docs


class PDFOcrReader(FileBaseReader):
    """PDF parser."""
    system_requirements = ['tesseract-ocr']
    requirements = ['PyMuPDF', 'pillow', 'pytesseract']

    def __init__(self, file: Path, single_text_per_document: bool = True, page_separator: str = '\n',
                 **load_kwargs):
        super().__init__(file, single_text_per_document, page_separator)
        self.load_kwargs = load_kwargs
        self.file = file

    def load_file(self, file: Path) -> List[Document]:
        import fitz
        from PIL import Image
        import tempfile
        from pytesseract import pytesseract
        docs = []
        pdf_document = fitz.open(file)
        for page_number in range(pdf_document.page_count):
            page = pdf_document[page_number]
            page_label = page.get_label()
            pix = page.get_pixmap()
            image = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
            with tempfile.TemporaryDirectory() as td:
                img_tmp = os.path.join(td, 'recdp_tmp.png')
                image.save(img_tmp)
                page_text = pytesseract.image_to_string(image)
                metadata = {"page_label": page_label, "source": str(file)}
                docs.append(Document(text=page_text, metadata=metadata))
        pdf_document.close()
        return docs


class DocxReader(FileBaseReader):
    """Docx parser."""
    system_requirements = []
    requirements = ['python-docx']

    def __init__(self, file: Path, single_text_per_document: bool = True, page_separator: str = '\n'):
        super().__init__(file, single_text_per_document, page_separator)

    def load_file(self, file: Path) -> List[Document]:
        """Parse file."""
        import docx
        document = docx.Document(str(file))

        # read in each paragraph in file
        return [Document(text=p.text, metadata={"source": str(file)}) for p in document.paragraphs]


class ImageReader(FileBaseReader):
    """Image parser.

    Extract text from images using pytesseract.

    """
    system_requirements = ['tesseract-ocr']
    requirements = ['pillow', 'pytesseract']

    def __init__(
            self,
            file: Path,
            single_text_per_document: bool = True,
            page_separator: str = '\n',
            keep_image: bool = False,
    ):
        super().__init__(file, single_text_per_document, page_separator)
        self._keep_image = keep_image

    def load_file(self, file: Path) -> List[Document]:
        """Parse file."""
        from PIL import Image
        from pytesseract import pytesseract
        # load document image
        image = Image.open(file)
        if image.mode != "RGB":
            image = image.convert("RGB")

        # Parse image into text
        text_str = pytesseract.image_to_string(image)

        return [
            Document(text=text_str, metadata={"source": str(file)})
        ]


class AudioReader(FileBaseReader):
    system_requirements = ['ffmpeg']
    requirements = ['openai-whisper']

    def __init__(
            self,
            file: Path,
            single_text_per_document: bool = True,
            page_separator: str = '\n',
            model='small',
    ):
        super().__init__(file, single_text_per_document, page_separator)
        import whisper
        self.model = whisper.load_model(model)

    def transcribe(self, file):
        return self.model.transcribe(file)

    def load_file(self, file: Path) -> List[Document]:
        file = str(file)
        result = self.transcribe(file)
        return [
            Document(text=result['text'], metadata={"source": file, 'language': result['language']})
        ]


_default_file_readers: Dict[str, Type[FileBaseReader]] = {
    ".pdf": PDFReader,
    ".docx": DocxReader,
    ".jpg": ImageReader,
    ".jpeg": ImageReader,
    ".png": ImageReader,
    ".mp3": AudioReader,
    ".wav": AudioReader,
    ".flac": AudioReader,
}


def read_from_directory(input_dirs: Optional[Union[str, List[str]]] = None,
                        input_files: Optional[list[str]] = None,
                        glob: str = "**/[!.]*",
                        recursive: bool = False,
                        exclude: Optional[List] = None,
                        exclude_hidden: bool = True,
                        max_concurrency: Optional[int] = None,
                        loaders: Optional[dict[str, Callable[[Path], List[Document]]]] = None,
                        required_exts: Optional[List[str]] = None,
                        pdf_ocr: bool = False,
                        ):
    """
   Loads documents from a directory or a list of files.

   Args:
       input_dirs: The input directory.
       input_files: A list of input files.
       glob: A glob pattern to match files.
       recursive: Whether to recursively search the input directory.
       exclude: A list of file patterns to exclude from loading.
       exclude_hidden: Whether to exclude hidden files from loading.
       max_concurrency: The maximum number of concurrent threads to use.
       loaders: Text Loader user what to use.
       required_exts: A list of file extensions that are required for documents.
                      default extensions are [.pdf, .docx, .jpeg, .jpg, .png]
       pdf_ocr: Whether to use ocr to load pdf.
   """
    if pdf_ocr:
        _default_file_readers['.pdf'] = PDFOcrReader

    def read_file(file_to_load, par):
        try:
            # use customer loader first if possible
            if loaders and file_to_load.suffix in loaders:
                document_loader: Callable = loaders[file_to_load.suffix]
                assert isinstance(document_loader, Callable)
                return document_loader(file_to_load)

            # load document with default registered file loader
            if file_to_load.suffix in _default_file_readers:
                document_loader_cls: Type[FileBaseReader] = _default_file_readers[file_to_load.suffix]
                loader = document_loader_cls(file_to_load)
                return loader.load()

            logger.warning(f"return empty documents, no loader registered for file type '{file_to_load.suffix}'")
            return []
        finally:
            par.update(1)

    if input_files:
        files_to_read = [Path(f) for f in input_files if os.path.isfile(f)]
    else:
        input_dirs = [input_dirs] if isinstance(input_dirs, str) else input_dirs
        from pyrecdp.core.path_utils import get_files
        files_to_read = [
            file_to_read
            for input_dir in input_dirs
            for file_to_read in get_files(
                input_dir,
                glob=glob,
                exclude=exclude,
                exclude_hidden=exclude_hidden,
                recursive=recursive,
                required_exts=required_exts,
            )
        ]

    files_to_read = list(set(files_to_read))

    if len(files_to_read) == 0:
        return []

    def install_reader_requirements():
        file_exts = set()
        for input_file in files_to_read:
            file_exts.add(input_file.suffix)

        for file_ext in file_exts:
            if file_ext in _default_file_readers:
                _default_file_readers[file_ext].setup()

    install_reader_requirements()

    from tqdm import tqdm
    from concurrent.futures import ThreadPoolExecutor
    pbar = tqdm(total=len(files_to_read))

    try:
        with ThreadPoolExecutor(max_concurrency) as executor:
            return [
                {'text': d.text, 'metadata': d.metadata}
                for docs in executor.map(lambda f: read_file(f, pbar), files_to_read)
                for d in docs
            ]
    finally:
        pbar.close()


def read_from_url(url: str, text_to_markdown: bool,
                  max_depth: Optional[int] = 1,
                  use_async: Optional[bool] = None,
                  extractor: Optional[Callable[[str], str]] = None,
                  metadata_extractor: Optional[Callable[[str, str], str]] = None,
                  exclude_dirs: Optional[Sequence[str]] = (),
                  timeout: Optional[int] = 10,
                  prevent_outside: bool = True,
                  link_regex: Union[str, re.Pattern, None] = None,
                  headers: Optional[dict] = None,
                  check_response_status: bool = False, ) -> List[dict[str, Any]]:
    """load documents from a url.
        Args:
            url: The URL to crawl.
            text_to_markdown: Whether to page as markdown text.
            max_depth: The max depth of the recursive loading.
            use_async: Whether to use asynchronous loading.
                If True, this function will not be lazy, but it will still work in the
                expected way, just not lazy.
            extractor: A function to extract document contents from raw html.
                When extract function returns an empty string, the document is
                ignored.
            metadata_extractor: A function to extract metadata from raw html and the
                source url (args in that order). Default extractor will attempt
                to use BeautifulSoup4 to extract the title, description and language
                of the page.
            exclude_dirs: A list of subdirectories to exclude.
            timeout: The timeout for the requests, in the unit of seconds. If None then
                connection will not timeout.
            prevent_outside: If True, prevent loading from urls which are not children
                of the root url.
            link_regex: Regex for extracting sub-links from the raw html of a web page.
            check_response_status: If True, check HTTP response status and skip
                URLs with error responses (400-599).
        """

    from langchain.document_loaders import RecursiveUrlLoader
    if not extractor:
        if text_to_markdown:
            def extractor_with_markdownify(x):
                import markdownify
                return markdownify.markdownify(x)

            extractor = extractor_with_markdownify
        elif extractor is None:
            def extractor_with_bs4(x):
                from bs4 import BeautifulSoup
                return BeautifulSoup(x, "html.parser").text

            extractor = extractor_with_bs4

    loader = RecursiveUrlLoader(
        url,
        extractor=extractor,
        max_depth=max_depth,
        use_async=use_async,
        metadata_extractor=metadata_extractor,
        exclude_dirs=exclude_dirs,
        timeout=timeout,
        prevent_outside=prevent_outside,
        link_regex=link_regex,
        headers=headers,
        check_response_status=check_response_status,
    )
    return [{'text': doc.page_content, 'metadata': doc.metadata} for doc in loader.load()]


def read_from_langchain(loader: str, loader_kwargs: Optional[dict[str, Any]] = None) -> List[dict[str, Any]]:
    """load documents using langchain document loader"""
    from pyrecdp.core.class_utils import new_instance
    from langchain.document_loaders.base import BaseLoader
    loader_kwargs = loader_kwargs or {}
    loader = new_instance("langchain.document_loaders", loader, **loader_kwargs)
    assert isinstance(loader, BaseLoader)
    return [{'text': doc.page_content, 'metadata': doc.metadata} for doc in loader.load()]


def transcribe_youtube_video(url: Union[str, List[str]], save_dir: Optional[str] = None, model_name: Optional[str] = None):
    """Load from YouTube transcript (if available) or audio.

    Args:
        url (Union[str, List[str]]): YouTube URL or list of URLs.
        save_dir (Optional[str]): Directory to save the audio files.
        model_name (Optional[str]): Model name for the transcription service.

    Returns:
        List[Dict]: List of documents with text and metadata.
    """
    urls = [url] if isinstance(url, str) else url
    
    # Try to load transcript from YouTube
    from langchain_community.document_loaders import YoutubeLoader
    pending_urls = []
    doc_list = []
    
    # Try to load transcript from YouTube
    for url in urls:
        try:
            loader = YoutubeLoader.from_youtube_url(url, add_video_info=True)
            docs = loader.load()
            if not docs:
                logger.info(f'No transcript available for URL: {url}. Defaulting to audio transcription.')
                pending_urls.append(url)
                continue
            for doc in docs:
                doc_list.append({'text': doc.page_content, 'metadata': doc.metadata})
        except Exception as e:
            logger.warning(f'Warning: Failed to load transcript from URL {url}. This may be due to language mismatch. Defaulting to audio transcription.')
            pending_urls.append(url)
    if len(pending_urls) == 0:
        return doc_list
    
    # If transcripts are not available, use Whisper for audio to text conversion
    import tempfile
    import shutil
    use_temp_dir = save_dir is None
    if use_temp_dir:
        save_dir = tempfile.mkdtemp()
    else:
        if not os.path.exists(save_dir):
            os.makedirs(save_dir)
    try:
        from langchain.document_loaders.blob_loaders.youtube_audio import YoutubeAudioLoader
        loader = YoutubeAudioLoader(urls, save_dir)
        audio_paths = {}
        for url, blob in zip(urls[::-1], loader.yield_blobs()):
            audio_paths[url] = str(blob.path)
        import whisper
        model = whisper.load_model(model_name)
        from pytube import Youtube
        audio_paths = {}
        for url in urls:
            video = YouTube(url)
            audio = video.streams.filter(only_audio=True, file_extension='mp4').first()
            audio_path = audio.download(output_path=save_dir)
            transcribe_result = model.transcribe(audio_path)
            doc_list.append({'text': transcribe_result['text'], 'metadata': {'source': url, 'language': transcribe_result['language']}})
    finally:
        if use_temp_dir:
            shutil.rmtree(save_dir)

    return doc_list
